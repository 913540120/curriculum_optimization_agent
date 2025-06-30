"""
培养方案文档解析器

支持解析PDF、Word、Excel等格式的培养方案文档，
提取课程信息、学分分配、技能映射等结构化数据
"""

import os
import re
import pandas as pd
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import asdict

# 文档解析相关
import PyPDF2
from docx import Document as DocxDocument
from openpyxl import load_workbook

# 项目模块
from ..models import CurriculumStructure, Course, CourseCategory, PracticalTraining
from openai import OpenAI
import json

class DocumentParser:
    """培养方案文档解析器"""
    
    def __init__(self, openai_api_key: str):
        """初始化解析器"""
        self.openai_client = OpenAI(
            base_url='https://api.siliconflow.cn/v1',
            api_key=openai_api_key
        )
        
        # 支持的文件格式
        self.supported_formats = ['.pdf', '.docx', '.xlsx', '.xls', '.txt']
        
        # 关键词模式匹配
        self.course_patterns = {
            'course_code': r'[A-Z]{2,4}\d{3,4}',  # 课程代码模式
            'credits': r'(\d+(?:\.\d+)?)\s*学分',   # 学分模式
            'hours': r'(\d+)\s*学时',              # 学时模式
            'semester': r'第?(\d+)学期',           # 学期模式
        }
        
        print("📄 文档解析器已初始化，支持格式:", self.supported_formats)

    def parse_document(self, file_path: str, major_name: str = "") -> CurriculumStructure:
        """
        解析培养方案文档
        
        Args:
            file_path: 文档路径
            major_name: 专业名称（可选）
            
        Returns:
            CurriculumStructure: 结构化的培养方案数据
        """
        print(f"🔍 开始解析文档: {file_path}")
        
        # 检查文件存在性和格式
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        if file_ext not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {file_ext}")
        
        # 根据文件类型解析
        try:
            if file_ext == '.pdf':
                raw_text = self._parse_pdf(file_path)
            elif file_ext == '.docx':
                raw_text = self._parse_docx(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                raw_text = self._parse_excel(file_path)
            else:  # .txt
                raw_text = self._parse_text(file_path)
            
            print(f"✅ 文档解析完成，提取文本长度: {len(raw_text)} 字符")
            
            # 使用AI进行结构化解析
            curriculum = self._ai_structured_parsing(raw_text, major_name)
            
            print("🎯 培养方案结构化解析完成")
            return curriculum
            
        except Exception as e:
            print(f"❌ 文档解析失败: {str(e)}")
            # 返回默认结构
            return self._create_default_curriculum(major_name)

    def _parse_pdf(self, file_path: str) -> str:
        """解析PDF文件"""
        print("📖 正在解析PDF文件...")
        text = ""
        
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    text += f"\n--- 第{page_num + 1}页 ---\n{page_text}"
        except Exception as e:
            print(f"⚠️ PDF解析出错: {e}")
            text = f"PDF解析失败: {str(e)}"
        
        return text

    def _parse_docx(self, file_path: str) -> str:
        """解析Word文档"""
        print("📝 正在解析Word文档...")
        text = ""
        
        try:
            doc = DocxDocument(file_path)
            
            # 提取段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # 提取表格内容
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                tables.append("\n".join(table_data))
            
            # 合并文本
            text = "\n".join(paragraphs)
            if tables:
                text += "\n\n=== 表格内容 ===\n" + "\n\n".join(tables)
                
        except Exception as e:
            print(f"⚠️ Word文档解析出错: {e}")
            text = f"Word文档解析失败: {str(e)}"
        
        return text

    def _parse_excel(self, file_path: str) -> str:
        """解析Excel文件"""
        print("📊 正在解析Excel文件...")
        text = ""
        
        try:
            # 使用pandas读取所有工作表
            xl_file = pd.ExcelFile(file_path)
            
            for sheet_name in xl_file.sheet_names:
                print(f"  📋 处理工作表: {sheet_name}")
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # 清理数据
                df = df.dropna(how='all').fillna('')
                
                # 转换为文本格式
                sheet_text = f"\n=== 工作表: {sheet_name} ===\n"
                sheet_text += df.to_string(index=False, max_cols=10)
                text += sheet_text + "\n"
                
        except Exception as e:
            print(f"⚠️ Excel解析出错: {e}")
            text = f"Excel解析失败: {str(e)}"
        
        return text

    def _parse_text(self, file_path: str) -> str:
        """解析纯文本文件"""
        print("📄 正在解析文本文件...")
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as file:
                    return file.read()
            except Exception as e:
                print(f"⚠️ 文本文件解析出错: {e}")
                return f"文本文件解析失败: {str(e)}"

    def _ai_structured_parsing(self, raw_text: str, major_name: str) -> CurriculumStructure:
        """使用AI进行结构化解析"""
        print("🤖 正在进行AI结构化解析...")
        
        system_prompt = """
        你是一个专业的培养方案解析专家。你的任务是从提供的培养方案文档中提取结构化信息。

        请从文档中提取以下信息：
        1. 基本信息：专业名称、学制、学位类型等
        2. 课程信息：课程代码、名称、学分、学时、开设学期、课程类别
        3. 学分分配：各类课程的学分分配情况
        4. 实践环节：实习、实训、毕业设计等
        5. 技能映射：课程对应培养的技能
        6. 毕业要求：总学分、各类学分要求等

        请严格按照以下JSON格式输出，确保所有字段都存在：
        """
        
        json_schema = """
        {
            "basic_info": {
                "major_name": "专业名称",
                "degree_type": "学位类型",
                "duration": "学制",
                "total_credits": "总学分"
            },
            "courses": [
                {
                    "code": "课程代码",
                    "name": "课程名称", 
                    "credits": 学分数值,
                    "hours": 学时数值,
                    "category": "课程类别",
                    "semester": 学期数值,
                    "prerequisites": ["先修课程"],
                    "description": "课程描述",
                    "skills": ["培养技能"]
                }
            ],
            "credit_distribution": {
                "通识教育": 学分数值,
                "专业基础": 学分数值,
                "专业核心": 学分数值,
                "专业选修": 学分数值,
                "实践教学": 学分数值
            },
            "practical_training": [
                {
                    "name": "实践环节名称",
                    "type": "实践类型",
                    "duration": 持续周数,
                    "credits": 学分数值,
                    "semester": 学期数值,
                    "objectives": ["实践目标"]
                }
            ],
            "skill_mapping": {
                "课程名称": ["技能1", "技能2"]
            },
            "graduation_requirements": {
                "total_credits": 总学分要求,
                "min_gpa": 最低GPA要求,
                "special_requirements": ["特殊要求"]
            }
        }
        """
        
        user_prompt = f"""
        专业名称（如果提供）: {major_name}
        
        培养方案文档内容：
        {raw_text[:8000]}  # 限制长度避免token超限
        
        请根据上述文档内容，提取结构化信息并按照指定的JSON格式输出。
        如果某些信息在文档中没有明确提及，请根据常见的培养方案模式进行合理推断。
        """
        
        try:
            print("🔄 调用AI模型进行结构化解析...")
            response = self.openai_client.chat.completions.create(
                model="deepseek-ai/DeepSeek-R1",
                messages=[
                    {"role": "system", "content": system_prompt + json_schema},
                    {"role": "user", "content": user_prompt}
                ],
                response_format={"type": "json_object"},
                timeout=120
            )
            
            # 解析AI响应
            ai_result = json.loads(response.choices[0].message.content)
            print("✅ AI结构化解析完成")
            
            # 转换为标准格式
            return self._convert_ai_result_to_curriculum(ai_result, major_name)
            
        except Exception as e:
            print(f"❌ AI解析失败: {e}")
            # 使用基础解析作为后备
            return self._basic_text_parsing(raw_text, major_name)

    def _convert_ai_result_to_curriculum(self, ai_result: Dict, major_name: str) -> CurriculumStructure:
        """将AI解析结果转换为标准的培养方案结构"""
        try:
            # 基本信息
            basic_info = ai_result.get("basic_info", {})
            if not basic_info.get("major_name") and major_name:
                basic_info["major_name"] = major_name
            
            # 课程信息转换
            courses = []
            for course_data in ai_result.get("courses", []):
                try:
                    # 确保必要字段存在
                    course_dict = {
                        "code": course_data.get("code", "UNKNOWN"),
                        "name": course_data.get("name", "未知课程"),
                        "credits": float(course_data.get("credits", 0)),
                        "hours": int(course_data.get("hours", 0)),
                        "category": course_data.get("category", "专业基础"),
                        "semester": int(course_data.get("semester", 1)),
                        "prerequisites": course_data.get("prerequisites", []),
                        "description": course_data.get("description", ""),
                        "skills": course_data.get("skills", [])
                    }
                    courses.append(course_dict)
                except (ValueError, TypeError) as e:
                    print(f"⚠️ 课程数据转换错误: {e}")
                    continue
            
            # 构建课程分类映射
            course_categories = {}
            for course in courses:
                category = course["category"]
                if category not in course_categories:
                    course_categories[category] = []
                course_categories[category].append(course["name"])
            
            # 构建培养方案结构
            curriculum = CurriculumStructure(
                basic_info=basic_info,
                courses=courses,
                course_categories=course_categories,
                credit_distribution=ai_result.get("credit_distribution", {}),
                prerequisite_relations={},  # 后续可以从课程数据中构建
                skill_mapping=ai_result.get("skill_mapping", {}),
                learning_outcomes=[],
                competency_framework={},
                practical_training=ai_result.get("practical_training", []),
                lab_experiments=[],
                graduation_project={},
                assessment_methods={},
                graduation_requirements=ai_result.get("graduation_requirements", {}),
                quality_standards=[]
            )
            
            return curriculum
            
        except Exception as e:
            print(f"❌ 数据转换失败: {e}")
            return self._create_default_curriculum(major_name)

    def _basic_text_parsing(self, text: str, major_name: str) -> CurriculumStructure:
        """基础文本解析（后备方案）"""
        print("🔧 使用基础文本解析...")
        
        # 简单的关键词提取
        courses = []
        
        # 查找可能的课程信息
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 尝试匹配课程代码
            code_match = re.search(self.course_patterns['course_code'], line)
            credit_match = re.search(self.course_patterns['credits'], line)
            
            if code_match or credit_match:
                # 提取课程信息
                course_code = code_match.group() if code_match else "UNKNOWN"
                credits = float(credit_match.group(1)) if credit_match else 2.0
                
                # 简化的课程名称提取（取第一个非代码部分）
                parts = line.split()
                course_name = "未知课程"
                for part in parts:
                    if not re.match(r'[A-Z]{2,4}\d{3,4}', part) and len(part) > 2:
                        course_name = part
                        break
                
                course_dict = {
                    "code": course_code,
                    "name": course_name,
                    "credits": credits,
                    "hours": int(credits * 16),  # 估算学时
                    "category": "专业基础",
                    "semester": 1,
                    "prerequisites": [],
                    "description": "",
                    "skills": []
                }
                courses.append(course_dict)
        
        # 如果没有找到课程，创建示例课程
        if not courses:
            courses = [
                {
                    "code": "CS101",
                    "name": "程序设计基础",
                    "credits": 3.0,
                    "hours": 48,
                    "category": "专业基础",
                    "semester": 1,
                    "prerequisites": [],
                    "description": "基础编程课程",
                    "skills": ["编程基础", "逻辑思维"]
                }
            ]
        
        return CurriculumStructure(
            basic_info={"major_name": major_name or "未知专业"},
            courses=courses,
            course_categories={"专业基础": [c["name"] for c in courses]},
            credit_distribution={"专业基础": sum(c["credits"] for c in courses)},
            prerequisite_relations={},
            skill_mapping={},
            learning_outcomes=[],
            competency_framework={},
            practical_training=[],
            lab_experiments=[],
            graduation_project={},
            assessment_methods={},
            graduation_requirements={},
            quality_standards=[]
        )

    def _create_default_curriculum(self, major_name: str) -> CurriculumStructure:
        """创建默认的培养方案结构"""
        print("🔧 创建默认培养方案结构...")
        
        default_courses = [
            {
                "code": "GE101",
                "name": "大学英语",
                "credits": 4.0,
                "hours": 64,
                "category": "通识教育",
                "semester": 1,
                "prerequisites": [],
                "description": "基础英语课程",
                "skills": ["英语交流", "跨文化沟通"]
            },
            {
                "code": "CS101",
                "name": "程序设计基础",
                "credits": 3.0,
                "hours": 48,
                "category": "专业基础",
                "semester": 1,
                "prerequisites": [],
                "description": "编程入门课程",
                "skills": ["编程基础", "算法思维"]
            }
        ]
        
        return CurriculumStructure(
            basic_info={
                "major_name": major_name or "计算机科学与技术",
                "degree_type": "工学学士",
                "duration": "四年",
                "total_credits": "160"
            },
            courses=default_courses,
            course_categories={
                "通识教育": ["大学英语"],
                "专业基础": ["程序设计基础"]
            },
            credit_distribution={
                "通识教育": 4.0,
                "专业基础": 3.0
            },
            prerequisite_relations={},
            skill_mapping={
                "大学英语": ["英语交流", "跨文化沟通"],
                "程序设计基础": ["编程基础", "算法思维"]
            },
            learning_outcomes=[],
            competency_framework={},
            practical_training=[],
            lab_experiments=[],
            graduation_project={},
            assessment_methods={},
            graduation_requirements={"total_credits": 160},
            quality_standards=[]
        )

    def validate_curriculum(self, curriculum: CurriculumStructure) -> Tuple[bool, List[str]]:
        """验证培养方案的完整性和合理性"""
        errors = []
        
        # 基本信息检查
        if not curriculum.get("basic_info", {}).get("major_name"):
            errors.append("缺少专业名称")
        
        # 课程信息检查
        courses = curriculum.get("courses", [])
        if not courses:
            errors.append("没有找到课程信息")
        
        total_credits = sum(course.get("credits", 0) for course in courses)
        if total_credits < 120:  # 一般本科要求最少120学分
            errors.append(f"总学分过少: {total_credits}")
        
        # 学分分配检查
        credit_dist = curriculum.get("credit_distribution", {})
        if not credit_dist:
            errors.append("缺少学分分配信息")
        
        is_valid = len(errors) == 0
        return is_valid, errors 